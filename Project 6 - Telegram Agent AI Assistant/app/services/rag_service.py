import os
import re
import csv
import sys
import subprocess
import logging
from typing import List, Dict, Any, Tuple

# --- Auto-install missing packages ---
def _ensure_package(import_name: str, pip_name: str = None):
    """Try importing a package; if missing, install it automatically."""
    pip_name = pip_name or import_name
    try:
        __import__(import_name)
    except ImportError:
        print(f"[auto-install] Installing missing package: {pip_name}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])

_ensure_package("pypdf")
_ensure_package("docx", "python-docx")
# --- End auto-install ---

import pypdf
import docx
from sqlalchemy.orm import Session

from app.config import settings
from app.models import Document, DocumentChunk

logger = logging.getLogger(__name__)

# Fallback ChromaDB / Embeddings importing
CHROMA_AVAILABLE = False
try:
    import chromadb
    from sentence_transformers import SentenceTransformer
    CHROMA_AVAILABLE = True
except ImportError:
    logger.warning("chromadb or sentence-transformers not installed. RAG will fall back to Database Token Matcher.")

class RAGService:
    def __init__(self):
        self.chroma_client = None
        self.chroma_collection = None
        self.transformer = None
        
        # Initialize real ChromaDB if available and not in Demo Mode
        if CHROMA_AVAILABLE and not settings.DEMO_MODE:
            try:
                self.chroma_client = chromadb.PersistentClient(path=settings.VECTOR_DB_PATH)
                self.chroma_collection = self.chroma_client.get_or_create_collection(
                    name="agent_kb",
                    metadata={"hnsw:space": "cosine"}
                )
                self.transformer = SentenceTransformer("all-MiniLM-L6-v2")
            except Exception as e:
                logger.error(f"Error initializing ChromaDB: {str(e)}. RAG fallback active.")
                self.chroma_client = None
                self.chroma_collection = None

    def read_pdf(self, file_path: str) -> str:
        """Extracts text from PDF."""
        text = ""
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
        return text

    def read_docx(self, file_path: str) -> str:
        """Extracts text from DOCX."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
        return text

    def read_txt(self, file_path: str) -> str:
        """Extracts text from TXT."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return ""

    def read_csv(self, file_path: str) -> str:
        """Extracts text summary/data from CSV."""
        rows_text = []
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                header = next(reader, None)
                if header:
                    rows_text.append("Columns: " + ", ".join(header))
                for i, row in enumerate(reader):
                    if i < 100:  # Cap lines to keep it manageable
                        rows_text.append(f"Row {i+1}: " + ", ".join(row))
                    else:
                        rows_text.append(f"... and {i} more rows")
                        break
        except Exception as e:
            logger.error(f"Error reading CSV {file_path}: {e}")
        return "\n".join(rows_text)

    def extract_text(self, file_path: str, file_type: str) -> str:
        """Decodes the document and returns flat text content."""
        file_type = file_type.lower()
        if file_type == "pdf":
            return self.read_pdf(file_path)
        elif file_type == "docx":
            return self.read_docx(file_path)
        elif file_type == "csv":
            return self.read_csv(file_path)
        else:
            return self.read_txt(file_path)

    def chunk_text(self, text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
        """Splits text into overlapping chunks."""
        words = text.split()
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        return chunks

    def ingest_document(self, db: Session, user_id: int, file_path: str, filename: str) -> int:
        """Parses a file, chunks it, stores in SQLite, and indexes in Vector DB."""
        file_type = filename.split(".")[-1].lower() if "." in filename else "txt"
        file_size = os.path.getsize(file_path)
        
        # 1. Extract text
        text = self.extract_text(file_path, file_type)
        chunks = self.chunk_text(text)
        
        # 2. Add document record to Database
        db_doc = Document(
            user_id=user_id,
            name=filename,
            path=file_path,
            file_type=file_type,
            file_size=file_size
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)

        # 3. Add chunks to database
        db_chunks = []
        for idx, chunk_content in enumerate(chunks):
            db_chunk = DocumentChunk(
                document_id=db_doc.id,
                chunk_index=idx,
                content=chunk_content,
                metadata_info={"page": (idx // 2) + 1, "filename": filename}
            )
            db.add(db_chunk)
            db_chunks.append(db_chunk)
        db.commit()

        # 4. Add to Vector DB (if active)
        if self.chroma_collection and self.transformer:
            try:
                ids = [f"doc_{db_doc.id}_chunk_{idx}" for idx in range(len(chunks))]
                embeddings = self.transformer.encode(chunks).tolist()
                metadatas = [{"document_id": db_doc.id, "filename": filename, "page": (idx // 2) + 1} for idx in range(len(chunks))]
                self.chroma_collection.add(
                    documents=chunks,
                    embeddings=embeddings,
                    metadatas=metadatas,
                    ids=ids
                )
            except Exception as e:
                logger.error(f"Error vector-indexing chunks in ChromaDB: {e}")
        
        return db_doc.id

    def retrieve(self, db: Session, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """Retrieves top_k context chunks from vector DB, with local SQL token fallback."""
        results = []
        
        # Method A: Try Chroma Vector Search
        if self.chroma_collection and self.transformer:
            try:
                query_embeddings = self.transformer.encode([query]).tolist()
                response = self.chroma_collection.query(
                    query_embeddings=query_embeddings,
                    n_results=top_k
                )
                if response and response["documents"] and len(response["documents"][0]) > 0:
                    for i in range(len(response["documents"][0])):
                        metadata = response["metadatas"][0][i]
                        results.append({
                            "content": response["documents"][0][i],
                            "filename": metadata.get("filename", "unknown"),
                            "page": metadata.get("page", 1),
                            "score": 1 - (response["distances"][0][i] if "distances" in response else 0.5)
                        })
                    return results
            except Exception as e:
                logger.error(f"Vector search failed: {e}. Falling back to token-matcher.")

        # Method B: Token Matcher Fallback
        # Extract keywords
        keywords = [w.lower() for w in re.findall(r'\b\w{4,}\b', query)]
        if not keywords:
            keywords = [w.lower() for w in query.split() if len(w) > 2]
            
        all_chunks = db.query(DocumentChunk).all()
        scored_chunks = []
        
        for chunk in all_chunks:
            score = 0
            content_lower = chunk.content.lower()
            for kw in keywords:
                if kw in content_lower:
                    score += 1
            if score > 0:
                scored_chunks.append((chunk, score))
                
        # Sort by match score
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        for chunk, score in scored_chunks[:top_k]:
            doc = db.query(Document).filter(Document.id == chunk.document_id).first()
            filename = doc.name if doc else "Document"
            page = chunk.metadata_info.get("page", 1) if chunk.metadata_info else 1
            results.append({
                "content": chunk.content,
                "filename": filename,
                "page": page,
                "score": float(score)
            })
            
        return results

    def delete_document(self, db: Session, document_id: int) -> bool:
        """Deletes a document and its chunks from SQL and vector stores."""
        doc = db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return False

        # Remove from vector store
        if self.chroma_collection:
            try:
                # Find number of chunks to build IDs
                chunks_count = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).count()
                ids = [f"doc_{document_id}_chunk_{idx}" for idx in range(chunks_count)]
                if ids:
                    self.chroma_collection.delete(ids=ids)
            except Exception as e:
                logger.error(f"ChromaDB delete error for doc {document_id}: {e}")

        # SQLite handles cascades to chunks, but let's delete explicitly to be safe
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
        db.delete(doc)
        db.commit()
        return True

rag_service = RAGService()
